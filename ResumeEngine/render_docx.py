"""render_docx.py — Generate a Word document from canonical data + tailored bullets.

Called by jd_analyzer.py Stage B after assemble_resume_md() writes the .md file.
Ported from output/corrected/_render_docx.py helper functions, adapted to work
with the canonical dict + tailored dict instead of hard-coded constants.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Section heading colour — navy blue, consistent with _render_docx.py
# ---------------------------------------------------------------------------
HEADING_COLOR = (0x1F, 0x3A, 0x5F)


# ---------------------------------------------------------------------------
# DOCX rendering helpers (ported from _render_docx.py)
# ---------------------------------------------------------------------------

def set_margins(doc: Document, top: float = 0.6, bottom: float = 0.6,
                left: float = 0.7, right: float = 0.7) -> None:
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_horizontal_line(paragraph) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(
    doc: Document,
    text: str,
    size: int = 12,
    bold: bool = True,
    color: tuple | None = None,
    space_before: int = 8,
    space_after: int = 2,
    border: bool = False,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if border:
        add_horizontal_line(p)
    return p


def add_para(
    doc: Document,
    text: str,
    size: int = 10,
    bold: bool = False,
    italic: bool = False,
    align=None,
    space_after: int = 2,
    indent: float | None = None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc: Document, text: str, size: int = 10, indent: float = 0.25):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    # Clear any auto-runs, then set the text
    for run in p.runs:
        run.font.size = Pt(size)
    run = p.add_run(text)
    run.font.size = Pt(size)
    # python-docx sometimes adds a blank run from the style; remove extra blanks
    for r in list(p.runs[:-1]):
        if not r.text:
            r._element.getparent().remove(r._element)
    return p


def add_labeled_line(doc: Document, label: str, body: str, size: int = 10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    lr.font.size = Pt(size)
    br = p.add_run(body)
    br.font.size = Pt(size)
    return p


def add_role_header(
    doc: Document,
    company: str,
    title: str,
    dates: str,
    location: str,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{title} | {company} | {dates}")
    run.bold = True
    run.font.size = Pt(11)
    add_para(doc, location, size=9, italic=True, space_after=2)


# ---------------------------------------------------------------------------
# Section builders — canonical data
# ---------------------------------------------------------------------------

def build_header(doc: Document, contact: dict) -> None:
    """Centred name + contact line. No tagline — summary section follows immediately."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(contact.get("name", ""))
    run.bold = True
    run.font.size = Pt(18)

    contact_line = (
        f"{contact.get('location', '')}  |  {contact.get('phone', '')}  |  "
        f"{contact.get('email', '')}  |  {contact.get('linkedin', '')}"
    )
    add_para(doc, contact_line, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)


def build_patents_section(doc: Document, canonical: dict) -> None:
    granted = canonical.get("patents_granted", [])
    pending = canonical.get("patents_pending", [])
    count_label = f"{len(granted)} Granted | {len(pending)} Pending"
    add_heading(doc, f"PATENTS ({count_label})", size=12, color=HEADING_COLOR, border=True)

    if granted:
        add_para(doc, "Granted:", size=10, bold=True, space_after=2)
        for p in granted:
            # Format: "US X,XXX,XXX — Title (Company, Year)"
            line = f"{p.get('number', '')} — {p.get('title', '')} ({p.get('company', '')}, {p.get('year', '')})"
            add_bullet(doc, line, size=10)

    if pending:
        add_para(doc, "Pending:", size=10, bold=True, space_after=2)
        for p in pending:
            line = f"{p.get('number', '')} — {p.get('title', '')} ({p.get('company', '')}, {p.get('year', '')})"
            add_bullet(doc, line, size=10)


def build_awards_section(doc: Document, canonical: dict) -> None:
    add_heading(doc, "AWARDS", size=12, color=HEADING_COLOR, border=True)
    for a in canonical.get("awards", []):
        line = f"{a.get('title', '')} ({a.get('year', '')}) — {a.get('subtitle', '')}"
        add_bullet(doc, line, size=10)


def build_affiliations_section(doc: Document, canonical: dict) -> None:
    affiliations = canonical.get("affiliations", [])
    if not affiliations:
        return
    add_heading(doc, "PROFESSIONAL AFFILIATIONS", size=12, color=HEADING_COLOR, border=True)
    for aff in affiliations:
        add_bullet(doc, aff, size=10)


def build_education_section(doc: Document, canonical: dict) -> None:
    add_heading(doc, "EDUCATION", size=12, color=HEADING_COLOR, border=True)
    edu = canonical.get("education", {})
    if edu.get("degree"):
        line = f"{edu['degree']} | {edu.get('school', '')} | {edu.get('dates', '')}"
    else:
        line = f"{edu.get('school', '')} | {edu.get('dates', '')}"
    add_para(doc, line, size=10)

    certs = canonical.get("certifications", [])
    if certs:
        for cert in certs:
            add_bullet(doc, cert, size=10)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def render_resume_docx(
    output_path: Path,
    canonical: dict,
    tailored: dict,
    jd_analysis: dict,
    template_name: str,
) -> None:
    """Build a .docx resume from canonical + tailored content.

    Args:
        output_path:   Destination .docx path.
        canonical:     Loaded from _canonical.yaml via load_canonical().
        tailored:      Stage A JSON output — professional_summary, core_skills, roles.
        jd_analysis:   Parsed JD signals (used for SA vs. IC heading labels).
        template_name: Template key — "resume-v2-sa" or "resume-v1-ic".
    """
    doc = Document()
    set_margins(doc, top=0.5, bottom=0.5, left=0.7, right=0.7)

    # --- Header ---
    build_header(doc, canonical["contact"])

    # --- Professional Summary ---
    add_heading(doc, "PROFESSIONAL SUMMARY", size=12, color=HEADING_COLOR, border=True)
    summary = tailored.get("professional_summary", "")
    if summary:
        add_para(doc, summary, size=10, space_after=4)

    # --- Core Skills / Competencies ---
    is_sa = "sa" in template_name
    skills_heading = "CORE COMPETENCIES" if is_sa else "CORE TECHNICAL SKILLS"
    add_heading(doc, skills_heading, size=12, color=HEADING_COLOR, border=True)
    for group_name, skills_str in tailored.get("core_skills", {}).items():
        add_labeled_line(doc, group_name, skills_str, size=10)

    # --- Professional Experience ---
    add_heading(doc, "PROFESSIONAL EXPERIENCE", size=12, color=HEADING_COLOR, border=True)
    for role_key, role_info in canonical.get("roles", {}).items():
        add_role_header(
            doc,
            company=role_info.get("company", ""),
            title=role_info.get("title", ""),
            dates=role_info.get("dates", ""),
            location=role_info.get("location", ""),
        )
        role_data = tailored.get("roles", {}).get(role_key, {})

        # NBCU departure note (canonical framing)
        if role_key == "nbcu":
            departure = role_data.get("nbcu_departure_note")
            if departure:
                add_para(doc, departure, size=10, italic=True, space_after=2)

        for bullet_text in role_data.get("bullets", []):
            # Strip leading "- " if present
            clean = bullet_text.strip().lstrip("- ").strip()
            if clean:
                add_bullet(doc, clean, size=10)

    # --- Patents (always from canonical — ALL of them) ---
    build_patents_section(doc, canonical)

    # --- Awards (always from canonical — ALL of them) ---
    build_awards_section(doc, canonical)

    # --- Affiliations ---
    build_affiliations_section(doc, canonical)

    # --- Education ---
    build_education_section(doc, canonical)

    doc.save(str(output_path))
