#!/usr/bin/env python3
"""
Render corrected resumes (factually verified against V1/V2/MST ground truth) to .docx.
Run: python3 _render_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent

# ---------------------------------------------------------------------------
# Canonical ground-truth data — verified against Eric's 3 past submitted resumes
# (ERIC MANCHESTER_MST.docx, Resume V1, Resume V2)
# ---------------------------------------------------------------------------

CONTACT = {
    "name": "ERIC MANCHESTER",
    "location": "Rockville, MD",
    "phone": "(240) 481-6737",
    "email": "eric.manchester@gmail.com",
    "linkedin": "linkedin.com/in/eric-manchester",
}

# All 4 Emmys Eric has won (verified)
AWARDS = [
    "Outstanding Interactive Experience Emmy (2025) — NBCUniversal, Paris 2024 Olympics; individually named Software Developer",
    "Technology & Engineering Emmy (2023) — NBCUniversal, Key Plays AI/ML sports highlights platform",
    "Emmy Award (2014) — Time Warner Cable, TWCTV Application (OTT innovation)",
    "Emmy Award (2005) — AOL, Live 8 Concert (world-record concurrent live streaming)",
]

# Television Academy membership (missing from all generated resumes!)
AFFILIATIONS = [
    "Television Academy — Science and Technology Peer Group Member",
]

# All 10 patents (7 granted + 3 pending, plus NDA) — verified against MST resume
PATENTS_GRANTED = [
    "US 8028092 — Inserting Advertising Content (AOL) — Dynamic insertion of advertising content into media streams",
    "US 8762575 — Inserting Advertising Content (TWC) — Enhanced targeting techniques within streaming video",
    "US 8769151 — Adding Advertising Content to Media Content (TWC) — Synchronizing content insertion with media playback",
    "US 10375452 — Apparatus and Methods for Thumbnail Generation (NBCU)",
    "US 10667018 — Asynchronous Workflows (NBCU) — Workflow orchestration for parallel distributed media processing",
    "US 11310567 — Apparatus and Methods for Thumbnail Generation (TWC)",
    "US 12382148 — Apparatus and Methods for Thumbnail Generation (AOL)",
]

PATENTS_PENDING = [
    "US 20250203152 — Timed Metadata for Overlays (NBCU) — Foundation of Peacock Live Actions",
    "US 20240430496 — Dynamic In-Scene Secondary Content Insertion (NBCU)",
    "AI-Driven Advertising Personalization — Under NDA",
]

# Education — no degree (theater major, did not graduate)
EDUCATION = "University of Maryland, College Park | 1990 – 1994"

# Certifications — AWS M&E Foundations just expired; omit per user instruction
# (do NOT list fabricated AWS Solutions Architect or Media & Entertainment Technology)
CERTIFICATIONS: list[str] = []

# ---------------------------------------------------------------------------
# Canonical job history — titles verified against V1/V2/MST
# Each entry is the raw data; tailoring (bullet selection + emphasis) happens per JD.
# CORRECTION: Eric's TWC 2012-2017 role was "Lead architect and IC" — NOT "independent architect"
# per direct instruction 2026-04-15. Source context files carry the legacy wording.
# ---------------------------------------------------------------------------

ROLES = {
    "harmonic": {
        "title": "Solutions Architect",
        "company": "Harmonic",
        "dates": "2025 – Present",
        "location": "Remote",
    },
    "re_kreations": {
        "title": "Founder and Principal R&D Architect",
        "company": "RE Kreations LLC",
        "dates": "2017 – Present",
        "location": "Rockville, MD",
    },
    "nbcu": {
        "title": "Principal End-to-End Architect",
        "company": "NBCUniversal",
        "dates": "2021 – 2025",
        "location": "Remote",
    },
    "leidos": {
        "title": "Innovation and Emerging Technology Lead",
        "company": "Leidos",
        "dates": "2019 – 2020",
        "location": "Washington, DC",
    },
    "ooyala": {
        "title": "Sr. Solutions Architect",
        "company": "Ooyala",
        "dates": "2017 – 2018",
        "location": "Remote",
    },
    "twc_2012": {
        "title": "Principal Engineer, Video-On-Demand",
        "company": "Time Warner Cable",
        "dates": "2012 – 2017",
        "location": "Remote",
    },
    "twc_2006": {
        "title": "Manager, Digital Media Delivery",
        "company": "Time Warner Cable",
        "dates": "2006 – 2012",
        "location": "Sterling, VA",
    },
    "aol": {
        "title": "Sr. Live Events Producer",
        "company": "AOL",
        "dates": "2001 – 2006",
        "location": "Sterling, VA",
    },
}

# ---------------------------------------------------------------------------
# DOCX rendering helpers
# ---------------------------------------------------------------------------

def set_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_heading(doc, text, size=12, bold=True, color=None, space_before=8, space_after=2, border=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if border:
        add_horizontal_line(p)
    return p

def add_para(doc, text, size=10, bold=False, italic=False, align=None, space_after=2, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, size=10, indent=0.25):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    for run in p.runs:
        run.font.size = Pt(size)
    run = p.add_run('')
    if not p.runs:
        run = p.add_run(text)
    else:
        p.runs[0].text = text
    for r in p.runs:
        r.font.size = Pt(size)
    return p

def add_labeled_line(doc, label, body, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    lr.font.size = Pt(size)
    br = p.add_run(body)
    br.font.size = Pt(size)
    return p

def add_role_header(doc, role_key, title_override=None):
    r = ROLES[role_key]
    title = title_override or r["title"]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{title} | {r['company']} | {r['dates']}")
    run.bold = True
    run.font.size = Pt(11)
    add_para(doc, r["location"], size=9, italic=True, space_after=2)

# ---------------------------------------------------------------------------
# Resume body builders — one per JD, each tailored to target
# All 3 share the same canonical job history + patents + awards + affiliations + education
# ---------------------------------------------------------------------------

def build_header(doc, tagline):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(CONTACT["name"])
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tagline)
    run.font.size = Pt(11)
    run.italic = True

    contact_line = (
        f"{CONTACT['location']}  |  {CONTACT['phone']}  |  "
        f"{CONTACT['email']}  |  {CONTACT['linkedin']}"
    )
    add_para(doc, contact_line, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

def build_patents_section(doc):
    add_heading(doc, "PATENTS (7 Granted | 3 Pending)", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc, "Granted:", size=10, bold=True, space_after=2)
    for p_text in PATENTS_GRANTED:
        add_bullet(doc, p_text, size=10)
    add_para(doc, "Pending:", size=10, bold=True, space_after=2)
    for p_text in PATENTS_PENDING:
        add_bullet(doc, p_text, size=10)

def build_awards_section(doc):
    add_heading(doc, "AWARDS", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    for a in AWARDS:
        add_bullet(doc, a, size=10)

def build_affiliations_section(doc):
    add_heading(doc, "PROFESSIONAL AFFILIATIONS", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    for a in AFFILIATIONS:
        add_bullet(doc, a, size=10)

def build_education_section(doc):
    add_heading(doc, "EDUCATION", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc, EDUCATION, size=10)
    if CERTIFICATIONS:
        for c in CERTIFICATIONS:
            add_bullet(doc, c, size=10)


# ---------------------------------------------------------------------------
# RESUME 1: Advisor360 — Staff/Tech Lead, AI-native B2B SaaS, player-coach role
# ---------------------------------------------------------------------------

def build_advisor360():
    doc = Document()
    set_margins(doc)

    build_header(doc, "Principal Engineer  |  AI-Native Platform Architecture  |  Technical Lead")

    add_heading(doc, "PROFESSIONAL SUMMARY", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc,
        "Principal-level engineer and architect with 26 years building cloud-native, distributed systems at broadcast and enterprise scale. "
        "Current focus: AI-native platform design — agentic workflows, LLM orchestration, MCP integrations, and CI-integrated quality gates "
        "for human + AI co-authored codebases. Proven track record as player-coach: setting architecture standards, reviewing code and specs "
        "for both engineers and AI agents, and writing the reference implementations that teach systems how to build themselves. "
        "4 Engineering Emmys, 7 granted patents, 3 pending. Member, Television Academy.",
        size=10, space_after=4)

    add_heading(doc, "CORE COMPETENCIES", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc,
        "Agentic Development Systems  |  LLM Orchestration (Claude, Ollama, Mixtral)  |  MCP & Tool-Use Patterns  |  "
        "Distributed Systems  |  Real-Time Data Processing  |  Event-Driven Architecture  |  "
        "CI-Integrated Quality Gates  |  AI-Powered Test Automation  |  Technical Leadership",
        size=10, space_after=4)

    add_heading(doc, "PROFESSIONAL EXPERIENCE", size=12, color=(0x1F, 0x3A, 0x5F), border=True)

    # Harmonic (current)
    add_role_header(doc, "harmonic")
    for b in [
        "Architect cloud-native, event-driven systems for high-availability content delivery, designing real-time data processing pipelines handling billions of requests with sub-second latency",
        "Lead presales architecture for tier-1 enterprise accounts, driving adoption through RFI/RFP responses, proof-of-concept development, and technical strategy sessions",
        "Serve as technical evangelist bridging complex distributed system requirements with business strategy for multi-platform delivery",
    ]:
        add_bullet(doc, b)

    # RE Kreations
    add_role_header(doc, "re_kreations")
    for b in [
        "Architected local-first, air-gapped AI platform (multi-model LLM orchestration via Ollama + Claude API) with automated ingest, classification, and retrieval across 60K+ vector chunks and sub-100ms RAG latency — zero external network dependencies, zero-backdoor data custody",
        "Designed and shipped agentic workflows using CLAUDE.md context artifacts, custom slash commands, and MCP tool-use patterns that teach AI agents how to operate inside a defined system boundary — the exact pattern Advisor360 is building for its product teams",
        "Built CI-integrated quality gates for AI-generated output: automated acceptance tests, diagnostic trace stubs, and format-repair validation — establishing a repeatable compound loop where each shipped feature improves agent reliability",
        "Refactored and consolidated system scripts, eliminated legacy paths, and standardized prompt routing across all LLM output layers — reducing pipeline fragility and enabling repeatable, auditable automation at scale",
    ]:
        add_bullet(doc, b)

    # NBCU
    add_role_header(doc, "nbcu")
    for b in [
        "Defined and owned organization-wide technical architecture across NBCUniversal direct-to-consumer platform, setting enterprise standards for systems handling millions of concurrent users from content ingestion through playback",
        "Architected production RAG pipelines and GenAI systems generating millions of personalized video experiences; redesigned metadata model to be AI-first, enabling personalization at scale",
        "Created Emmy Award-winning AI/ML architecture (Key Plays) for automated event detection with real-time data processing and rich metadata generation",
        "Invented and shipped patent-pending server-side content overlay system (Live Actions, US 20250203152) delivering time-based synchronized experiences across millions of concurrent users",
        "Delivered ad innovation architecture driving 200%+ revenue increase through enhanced targeting; partnered with ad ops, sales, and architecture teams to integrate monetization into platform roadmaps",
        "Mentored cross-functional team members including a project manager who grew from onboarding into the lead PM driving ads innovation strategy and execution",
        "Note: Departed during organizational restructuring in early 2025.",
    ]:
        add_bullet(doc, b)

    # Leidos
    add_role_header(doc, "leidos")
    for b in [
        "Drove R&D initiatives for federal agency client in highly regulated domain, developing advanced technology concepts and presenting demonstrations to C-level executives",
        "Designed secure, scalable solutions for regulated environments while maintaining innovation velocity",
    ]:
        add_bullet(doc, b)

    # Ooyala
    add_role_header(doc, "ooyala")
    for b in [
        "Architected SaaS-based real-time data processing and content management solutions, partnering with C-suite executives as SME on distributed systems and data pipeline optimization",
        "Built fully automated workflows for large-scale data processing; platform served as full Media Asset Management solution",
    ]:
        add_bullet(doc, b)

    # TWC 2012-2017 — CORRECTED: removed "independent architect" wording per user direction
    add_role_header(doc, "twc_2012")
    for b in [
        "Lead architect and IC on TWCTV — established organization-wide technical standards for VOD across CableLabs, SCTE, and SMPTE specifications; designed multi-layered access control and DRM systems protecting multimillion-dollar premium content",
        "Centralized fragmented VOD architecture across all regions into a unified platform; led team of engineers through acquisition, processing, and distribution pipeline design; mentored engineers from NOC operations into architecture roles",
        "Built end-to-end data provenance and chain-of-custody systems tracking every content asset from source acquisition through distribution — ensuring complete auditability across a multi-stage workflow",
        "Led Emmy Award-winning TWCTV Application architecture (2014)",
    ]:
        add_bullet(doc, b)

    # TWC 2006-2012
    add_role_header(doc, "twc_2006")
    for b in [
        "Led engineering team architecting and deploying fully automated content capture and encoding workflows, driving catalog migration to SaaS platform",
        "Directed metadata enrichment and CMS automation initiatives; ad revenue grew 350%+ through dynamic insertion architecture",
    ]:
        add_bullet(doc, b)

    # AOL
    add_role_header(doc, "aol")
    for b in [
        "Architected and built industry-first all-digital Broadcast Operations Center; secured multi-million-dollar budget through executive presentations",
        "Delivered Emmy Award-winning Live 8 Concert (2005) from eight simultaneous global sites, setting world record for concurrent live streaming viewers",
    ]:
        add_bullet(doc, b)

    build_patents_section(doc)
    build_awards_section(doc)
    build_affiliations_section(doc)

    add_heading(doc, "TECHNICAL EXPERTISE", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_labeled_line(doc, "AI & Agentic Systems", "Claude Code, Cursor, MCP, custom slash commands, agent hooks, CLAUDE.md context artifacts, LLM orchestration, RAG pipelines, ASR, OCR, multi-model coordination")
    add_labeled_line(doc, "LLM Stack (hands-on)", "Claude (Sonnet/Haiku), Ollama, Mixtral, Llama, Mistral, Phi-3, Nomic Embed Text, Whisper; Anthropic SDK, OpenAI SDK, FAISS vector search")
    add_labeled_line(doc, "Platform & Distributed Systems", "Cloud-native microservices, distributed edge services, event-driven architecture, multi-tenant SaaS, REST APIs, real-time processing, high-availability design")
    add_labeled_line(doc, "Cloud & Infrastructure", "AWS, GCP, Azure, infrastructure automation, multi-CDN delivery, CI/CD pipelines")
    add_labeled_line(doc, "Test Automation", "E2E test pipeline design, CI-integrated quality gates, regression suites, diagnostic trace frameworks, AI-generated-code review")
    add_labeled_line(doc, "Languages & Development", "Python, TypeScript, Agile, API design, observability, structured output enforcement, audit logging")
    add_labeled_line(doc, "Standards & Protocols", "HLS, DASH, CMAF, SCTE-35, VAST, VPAID")

    build_education_section(doc)

    out = HERE / "Advisor360_TechLead.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# RESUME 2: Airbnb — Senior Backend Engineer, Media Foundation (media pipelines, ML infra)
# ---------------------------------------------------------------------------

def build_airbnb():
    doc = Document()
    set_margins(doc)

    build_header(doc, "Senior Backend Engineer  |  Media Processing Platforms  |  Distributed Systems at Scale")

    add_heading(doc, "PROFESSIONAL SUMMARY", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc,
        "Backend systems architect with 26 years building large-scale media processing platforms, distributed pipeline design, and "
        "ML/AI-enabled infrastructure. Designed end-to-end systems that ingest, process, store, and serve billions of media assets "
        "across OTT, broadcast, and streaming surfaces — with a record of enabling product teams to ship richer, smarter media experiences "
        "through self-serve platform capabilities. 4 Engineering Emmys, 7 granted patents (4 directly on media processing workflows), "
        "3 pending. Member, Television Academy.",
        size=10, space_after=4)

    add_heading(doc, "CORE COMPETENCIES", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc,
        "Media Processing Pipelines  |  Image/Video Transcoding  |  Metadata Extraction & Enrichment  |  "
        "ML/AI Workflow Infrastructure  |  Distributed Systems  |  Backend API Design  |  "
        "Event-Driven Architecture  |  Cost Optimization & Vendor Consolidation",
        size=10, space_after=4)

    add_heading(doc, "PROFESSIONAL EXPERIENCE", size=12, color=(0x1F, 0x3A, 0x5F), border=True)

    add_role_header(doc, "harmonic")
    for b in [
        "Design end-to-end cloud-native media platform architectures (VOS360/XOS) for tier-1 broadcasters and streaming operators — translating complex content delivery requirements into scalable backend and distribution solutions",
        "Architect SSAI and dynamic ad insertion workflows integrated with live and VOD media pipelines, enabling downstream product teams to build monetization experiences on shared media infrastructure",
        "Evaluate storage optimization and vendor consolidation strategies for customers operating multi-CDN, multi-format media delivery at scale",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "re_kreations")
    for b in [
        "Independently designed and operate a local-first AI/LLM orchestration platform using multi-model pipelines (Python, Ollama, Claude API), vector search (FAISS), and RAG architecture — handling 60K+ document chunks with sub-100ms retrieval latency",
        "Built end-to-end batch processing and metadata enrichment systems for media and knowledge assets, applying the same pipeline architecture patterns used at broadcast scale to a founder-led product environment",
        "Architected microservice APIs and data models supporting contracted client engagements across content management and media intelligence use cases",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "nbcu")
    for b in [
        "Architected Key Plays (Emmy Award, 2023) — a near real-time AI/ML highlight detection and delivery platform processing live video feeds through ML analysis pipelines, extracting moment metadata, and serving packaged highlights to millions of concurrent users across NFL, NBA, Premier League, and Olympics",
        "Designed the backend orchestration layer for ML/AI analysis workflows on media assets, enabling model-triggered analysis, structured metadata storage, and downstream API access for product teams building interactive sports surfaces — directly analogous to self-serve platform capabilities for ML/AI teams",
        "Engineered Olympic Daily Recap personalization (Emmy Award, 2025) generating 7M+ unique video variants via GenAI voice synthesis and asynchronous content assembly — batch-scale media processing and metadata-driven personalization at production grade",
        "Authored US Patent 10,667,018 (Asynchronous Workflows) covering event-driven task orchestration and state management patterns foundational to large-scale content processing",
        "Delivered Live Actions (patent pending, US 20250203152) — real-time metadata ingestion, event detection, and synchronized overlay delivery across streaming and broadcast",
        "Note: Departed during organizational restructuring in early 2025.",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "leidos")
    for b in [
        "Delivered technical consultation on distributed media and data platform architectures for government and enterprise clients",
        "Designed solution architectures integrating SaaS media processing platforms with government data workflows",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "ooyala")
    for b in [
        "Architected SaaS-based media lifecycle platform solutions for enterprise broadcasting and streaming clients — ingest, processing, and distribution workflows on top of Ooyala's media infrastructure",
        "Aligned C-suite and engineering stakeholders across client organizations to map complex media processing requirements to scalable backend capabilities",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "twc_2012")
    for b in [
        "Lead architect and IC on TWCTV VOD platform — centralized all content acquisition, processing, and distribution into a unified headend architecture serving millions of cable subscribers",
        "Architected just-in-time encryption DRM solutions for VOD content delivery, establishing media security and access control patterns scaled across OTT and set-top surfaces",
        "Led thumbnail generation patent work (US 11310567) — automated preview generation for media navigation at catalog scale",
        "Recognized with Emmy Award (2014) as key architect on TWCTV Application",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "twc_2006")
    for b in [
        "Led engineering team building fully automated catch-and-encoding workflows; transitioned VMS/CMS to SaaS platform",
        "Directed metadata enrichment and CMS automation; ad revenue grew 350%+ through dynamic insertion architecture",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "aol")
    for b in [
        "Architected industry's first all-digital Broadcast Operations Center for streaming/OTT content delivery, establishing real-time video ingestion, processing, and global distribution at a time when no industry playbook existed",
        "Delivered Live 8 Concert streaming (Emmy Award, 2005) — world record for concurrent streaming viewers across 8 global sites",
    ]:
        add_bullet(doc, b)

    build_patents_section(doc)
    build_awards_section(doc)
    build_affiliations_section(doc)

    add_heading(doc, "TECHNICAL EXPERTISE", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_labeled_line(doc, "Media Processing", "Image processing, video transcoding, thumbnail generation, metadata extraction, media ingestion pipelines, VOD architecture, HLS, DASH, CMAF, H.264, H.265")
    add_labeled_line(doc, "Backend Systems & APIs", "Python, REST API design, data modeling, microservices, asynchronous workflow orchestration, batch processing, high-throughput pipelines")
    add_labeled_line(doc, "ML/AI Workflow Infrastructure", "Near real-time ML model integration, AI-driven metadata enrichment, personalization pipelines, multi-model orchestration (Ollama, Claude, Mistral, Mixtral, Phi-3), vector search (FAISS), RAG")
    add_labeled_line(doc, "Cloud & Distributed Systems", "AWS, GCP, Azure, distributed system design, event-driven architecture, scalable platform design, multi-CDN operations")
    add_labeled_line(doc, "Storage & Data", "Content metadata management, DRM/encryption, database modeling, storage optimization, vendor consolidation, SaaS provider integration")
    add_labeled_line(doc, "Standards", "SCTE-35, VAST, VPAID, SMPTE, CableLabs")

    build_education_section(doc)

    out = HERE / "AirBNB_BackEnd_MediaIngest.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# RESUME 3: Akamai — Principal Technical Solutions Architect (presales, cloud/edge/media)
# ---------------------------------------------------------------------------

def build_akamai():
    doc = Document()
    set_margins(doc)

    build_header(doc, "Principal Solutions Architect  |  Customer-Facing Technical Leadership  |  Cloud-Native & Media Streaming")

    add_heading(doc, "PROFESSIONAL SUMMARY", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc,
        "Solutions architect with 26 years driving customer-facing technical engagements and architecting cloud-native, distributed media "
        "platforms at broadcast scale. Extensive presales experience for tier-1 enterprise accounts — translating complex business "
        "requirements into scalable technical solutions, leading RFI/RFP responses, developing proofs of concept, and partnering with "
        "C-suite executives as trusted technical advisor. Deep expertise in media streaming (live transcoding, VOD workflows, multi-CDN "
        "delivery), cloud-native microservices, and distributed edge services. 4 Engineering Emmys, 7 granted patents, 3 pending. "
        "Member, Television Academy.",
        size=10, space_after=4)

    add_heading(doc, "CORE COMPETENCIES", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_para(doc,
        "Customer-Facing Solution Architecture  |  Presales & Proof-of-Concept  |  RFI/RFP Response  |  "
        "Technical Evangelism  |  Executive-Level Communication  |  Media Streaming (Live/VOD)  |  "
        "Multi-CDN Delivery  |  Cloud-Native Microservices  |  Distributed Systems",
        size=10, space_after=4)

    add_heading(doc, "PROFESSIONAL EXPERIENCE", size=12, color=(0x1F, 0x3A, 0x5F), border=True)

    add_role_header(doc, "harmonic")
    for b in [
        "Lead customer-facing presales architecture for tier-1 enterprise broadcasters and streaming operators, driving cloud-native adoption through direct engagement — RFI/RFP responses, proof-of-concept development, technical presentations, and solution design sessions",
        "Architect cloud-native, event-driven systems for high-availability content delivery on VOS360/XOS, designing real-time data processing pipelines handling billions of requests with sub-second latency",
        "Serve as technical evangelist bridging complex distributed system requirements with business strategy, presenting solutions to C-level executives and technical stakeholders across multi-platform delivery environments",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "re_kreations")
    for b in [
        "Founded and operate R&D practice building AI/ML and media infrastructure solutions for contracted clients — operating autonomously across full product lifecycle from requirements through production",
        "Architected local-first LLM orchestration platform (Ollama + Claude API + FAISS) with 60K+ vector chunks and sub-100ms retrieval — emerging-tech credibility directly relevant to architecting next-generation edge-native applications",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "nbcu")
    for b in [
        "Architected Key Plays (Technology & Engineering Emmy, 2023) — near real-time AI/ML sports highlights platform serving millions of concurrent users on Peacock",
        "Designed Olympic Daily Recap (Outstanding Interactive Experience Emmy, 2025) — 7M+ personalized video variants for Paris 2024 Olympics; individually named Software Developer on Emmy",
        "Invented and shipped Live Actions (patent pending, US 20250203152) — real-time interactive overlays synchronized with live broadcast events at tens of millions of concurrent users",
        "Delivered ad innovation architecture driving 200%+ revenue increase through enhanced targeting and dynamic insertion",
        "Operated with full autonomy on portfolio of concurrent projects — partnered directly with product, business, and operations stakeholders to translate strategic objectives into technical solutions",
        "Note: Departed during organizational restructuring in early 2025.",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "leidos")
    for b in [
        "Drove R&D and emerging-technology initiatives for federal agency client, presenting technical demonstrations to C-level executives",
        "Created innovation pipeline for evaluating and integrating emerging technologies; mentored staff on transitioning concepts to production-ready systems",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "ooyala")
    for b in [
        "Architected cloud-based SaaS media lifecycle solutions for Fortune 500 media customers, partnering directly with C-suite executives as technical SME on distributed systems, cloud infrastructure, and data pipelines",
        "Engaged customers through technical discovery, solution design, proof-of-concept development, and deployment support across complex multi-cloud and hybrid environments",
        "Built fully automated workflows for large-scale data processing; platform served as full Media Asset Management solution while integrating with existing enterprise systems",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "twc_2012")
    for b in [
        "Lead architect and IC on TWCTV — established organization-wide technical standards for VOD across CableLabs, SCTE, and SMPTE specifications; designed multi-layered access control and DRM systems",
        "Centralized fragmented VOD architecture across all regions into a unified, CDN-distributed platform serving millions of cable subscribers; mentored engineers from NOC operations into architecture roles",
        "Acted as SME for stakeholders from engineers to C-suite on content acquisition, processing, CDN-based distribution, and playback — translating complex technical tradeoffs into executive-aligned decisions",
        "Emmy Award (2014) for TWCTV Application",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "twc_2006")
    for b in [
        "Led engineering team building fully automated catch-and-encoding workflows; transitioned VMS/CMS to SaaS platform",
        "Directed metadata enrichment and CMS automation; ad revenue grew 350%+ through dynamic insertion",
    ]:
        add_bullet(doc, b)

    add_role_header(doc, "aol")
    for b in [
        "Architected industry's first all-digital Broadcast Operations Center for streaming/OTT delivery — CDN-distributed infrastructure serving millions of concurrent viewers",
        "Delivered Live 8 Concert (Emmy Award, 2005) — world record for concurrent live streaming across 8 global sites",
    ]:
        add_bullet(doc, b)

    build_patents_section(doc)
    build_awards_section(doc)
    build_affiliations_section(doc)

    add_heading(doc, "TECHNICAL EXPERTISE", size=12, color=(0x1F, 0x3A, 0x5F), border=True)
    add_labeled_line(doc, "Solution Architecture", "Customer-facing technical leadership, presales, proof-of-concept delivery, RFI/RFP, executive communication, technical evangelism, strategic account engagement")
    add_labeled_line(doc, "Media Streaming", "Live transcoding, VOD workflows, HLS, DASH, CMAF, multi-DRM integration (DRM standards & CableLabs spec), video players, broadcast-grade ingest")
    add_labeled_line(doc, "CDN & Content Delivery", "Multi-CDN delivery strategy, origin/edge distribution, real-time video delivery at scale, content protection, SSAI/CSAI")
    add_labeled_line(doc, "Cloud & Distributed Systems", "Cloud-native microservices, distributed edge services, event-driven architecture, multi-tenant SaaS, REST APIs, real-time data processing, high-availability design — AWS, GCP, Azure")
    add_labeled_line(doc, "AI/ML (hands-on)", "GenAI, RAG pipelines, LLM orchestration (Claude, Ollama, Mixtral, Llama, Mistral, Phi-3), vector search (FAISS), AI-first metadata systems")
    add_labeled_line(doc, "Standards & Protocols", "HLS, DASH, CMAF, SCTE-35, VAST, VPAID, SMPTE, CableLabs")

    build_education_section(doc)

    out = HERE / "Akamai_PrinTechSolArch.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print("Rendering 3 corrected resumes...")
    p1 = build_advisor360()
    print(f"  Advisor360 → {p1}")
    p2 = build_airbnb()
    print(f"  Airbnb     → {p2}")
    p3 = build_akamai()
    print(f"  Akamai     → {p3}")
    print("Done.")
